from __future__ import annotations

import argparse
import cgi
import json
import mimetypes
import os
from http.server import SimpleHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from tempfile import NamedTemporaryFile
from urllib.parse import unquote

import pdfplumber
from docx import Document


ROOT = Path(__file__).resolve().parent
MAX_UPLOAD_BYTES = 20 * 1024 * 1024


class ResumeAgentHandler(SimpleHTTPRequestHandler):
    server_version = "ResumeAgent/1.0"

    def translate_path(self, path: str) -> str:
        path = unquote(path.split("?", 1)[0].split("#", 1)[0])
        parts = [part for part in path.split("/") if part and part not in {".", ".."}]
        resolved = ROOT.joinpath(*parts).resolve()
        if not str(resolved).startswith(str(ROOT)):
            return str(ROOT / "index.html")
        return str(resolved)

    def do_POST(self) -> None:
        if self.path != "/api/parse-resume":
            self.send_json({"ok": False, "error": "接口不存在"}, status=404)
            return

        content_length = int(self.headers.get("Content-Length", "0") or "0")
        if content_length <= 0:
            self.send_json({"ok": False, "error": "没有收到文件"}, status=400)
            return
        if content_length > MAX_UPLOAD_BYTES:
            self.send_json({"ok": False, "error": "文件超过 20MB，请先压缩或导出为文本"}, status=413)
            return

        content_type = self.headers.get("Content-Type", "")
        if "multipart/form-data" not in content_type:
            self.send_json({"ok": False, "error": "请使用文件上传表单"}, status=400)
            return

        form = cgi.FieldStorage(
            fp=self.rfile,
            headers=self.headers,
            environ={
                "REQUEST_METHOD": "POST",
                "CONTENT_TYPE": content_type,
                "CONTENT_LENGTH": str(content_length),
            },
        )
        field = form["resume"] if "resume" in form else None
        if field is None or not getattr(field, "filename", ""):
            self.send_json({"ok": False, "error": "没有找到 resume 文件字段"}, status=400)
            return

        filename = os.path.basename(field.filename)
        suffix = Path(filename).suffix.lower()
        data = field.file.read()

        try:
            text = extract_resume_text(data, suffix)
        except ValueError as exc:
            self.send_json({"ok": False, "error": str(exc)}, status=400)
            return
        except Exception as exc:
            self.send_json({"ok": False, "error": f"解析失败：{exc}"}, status=500)
            return

        if not text.strip():
            self.send_json({"ok": False, "error": "未提取到文本。扫描版 PDF 需要先 OCR 后再导入。"}, status=422)
            return

        self.send_json({"ok": True, "filename": filename, "text": normalize_text(text)})

    def end_headers(self) -> None:
        self.send_header("Access-Control-Allow-Origin", "*")
        super().end_headers()

    def send_json(self, payload: dict, status: int = 200) -> None:
        body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)


def extract_resume_text(data: bytes, suffix: str) -> str:
    if suffix in {".txt", ".md", ".text"}:
        return decode_text(data)
    if suffix == ".pdf":
        return extract_pdf_text(data)
    if suffix == ".docx":
        return extract_docx_text(data)
    if suffix == ".doc":
        raise ValueError("暂不支持旧版 .doc 二进制格式。请用 Word 另存为 .docx 或 PDF 后再导入。")
    raise ValueError("仅支持 PDF、DOCX、TXT 和 Markdown 简历。")


def decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def extract_pdf_text(data: bytes) -> str:
    with NamedTemporaryFile(suffix=".pdf", delete=False) as temp:
        temp.write(data)
        temp_path = temp.name
    try:
        chunks: list[str] = []
        with pdfplumber.open(temp_path) as pdf:
            for page in pdf.pages:
                chunks.append(page.extract_text(x_tolerance=1, y_tolerance=3) or "")
        return "\n".join(chunks)
    finally:
        try:
            os.remove(temp_path)
        except OSError:
            pass


def extract_docx_text(data: bytes) -> str:
    with NamedTemporaryFile(suffix=".docx", delete=False) as temp:
        temp.write(data)
        temp_path = temp.name
    try:
        document = Document(temp_path)
        chunks: list[str] = []
        chunks.extend(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    chunks.append(" | ".join(cells))
        return "\n".join(chunks)
    finally:
        try:
            os.remove(temp_path)
        except OSError:
            pass


def normalize_text(text: str) -> str:
    lines = [line.strip() for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    compacted: list[str] = []
    previous_blank = False
    for line in lines:
        is_blank = not line
        if is_blank and previous_blank:
            continue
        compacted.append(line)
        previous_blank = is_blank
    return "\n".join(compacted).strip()


def main() -> None:
    parser = argparse.ArgumentParser(description="Resume Agent local server")
    parser.add_argument("--host", default="0.0.0.0")
    parser.add_argument("--port", type=int, default=8767)
    args = parser.parse_args()

    mimetypes.add_type("text/javascript; charset=utf-8", ".js")
    mimetypes.add_type("text/css; charset=utf-8", ".css")
    os.chdir(ROOT)
    server = ThreadingHTTPServer((args.host, args.port), ResumeAgentHandler)
    print(f"Resume Agent server: http://{args.host}:{args.port}/index.html", flush=True)
    server.serve_forever()


if __name__ == "__main__":
    main()
